#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板分析器 — 审计附注模板 Word 文档表格结构分析工具
====================================================

功能:
  1. 读取 .docx 审计附注模板，提取所有表格的结构信息
  2. 推断表格名称（通过前文段落匹配）
  3. 检测表内小节标题（一、二、三… / 1．2．3．…）
  4. 标记合计/小计/总计行
  5. 标记 XX/XXX/XXXX 占位符单元格
  6. 输出可序列化的结构化字典

用法:
  python template_analyzer.py <模板路径.docx>
  python template_analyzer.py --json <模板路径.docx>   # JSON 格式输出
"""

import sys
import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Any

# ── 确保环境 ──────────────────────────────────────────────
try:
    from docx import Document
except ImportError:
    sys.exit("错误: 需要 python-docx 库。请运行: pip install python-docx")


# ============================================================
# 公共附注表名模式库
# ============================================================

# 标准财务报表附注表名（按资产负债表→利润表→现金流量表→附注明细顺序）
COMMON_NOTE_TABLE_NAMES = [
    # 资产类
    "货币资金",
    "交易性金融资产",
    "应收票据",
    "应收账款",
    "应收款项融资",
    "预付账款",
    "应收股利",
    "应收利息",
    "其他应收款",
    "存货",
    "合同资产",
    "持有待售资产",
    "一年内到期的非流动资产",
    "其他流动资产",
    "债权投资",
    "其他债权投资",
    "长期应收款",
    "长期股权投资",
    "其他权益工具投资",
    "其他非流动金融资产",
    "投资性房地产",
    "固定资产",
    "在建工程",
    "生产性生物资产",
    "油气资产",
    "使用权资产",
    "无形资产",
    "开发支出",
    "商誉",
    "长期待摊费用",
    "递延所得税资产",
    "其他非流动资产",
    # 负债类
    "短期借款",
    "应付票据",
    "应付账款",
    "预收账款",
    "合同负债",
    "应付职工薪酬",
    "应交税费",
    "应付利息",
    "应付股利",
    "其他应付款",
    "持有待售负债",
    "一年内到期的非流动负债",
    "其他流动负债",
    "长期借款",
    "应付债券",
    "长期应付款",
    "租赁负债",
    "预计负债",
    "递延收益",
    "递延所得税负债",
    "其他非流动负债",
    # 权益类
    "实收资本",
    "股本",
    "资本公积",
    "减：库存股",
    "其他综合收益",
    "专项储备",
    "盈余公积",
    "未分配利润",
    # 损益类
    "营业收入",
    "营业成本",
    "税金及附加",
    "销售费用",
    "管理费用",
    "研发费用",
    "财务费用",
    "其他收益",
    "投资收益",
    "公允价值变动收益",
    "信用减值损失",
    "资产减值损失",
    "资产处置收益",
    "营业外收入",
    "营业外支出",
    "所得税费用",
    # 现金流量表
    "现金流量表",
    "现金流量表补充资料",
    # 特殊
    "应收账款-账龄",
    "应收账款-组合",
    "其他应收款-账龄",
    "其他应收款-组合",
    "固定资产-明细",
    "无形资产-明细",
    "在建工程-明细",
    "长期待摊费用-明细",
    "应付账款-账龄",
    "成本费用管理情况",
    "营业外收支",
    "资产减值损失明细",
    "信用减值损失明细",
]

# 表名中常见的关键词 → 用于从不完整的段落文本中推断表名
TABLE_NAME_KEYWORDS = {
    "货币资金": ["货币资金", "库存现金", "银行存款", "其他货币资金"],
    "应收账款": ["应收账款", "应收账"],
    "其他应收款": ["其他应收款", "其他应收"],
    "预付账款": ["预付账款", "预付"],
    "固定资产": ["固定资产", "固定"],
    "在建工程": ["在建工程", "在建"],
    "无形资产": ["无形资产", "无形"],
    "长期待摊费用": ["长期待摊费用", "长期待摊"],
    "应付账款": ["应付账款", "应付"],
    "应付职工薪酬": ["应付职工薪酬", "应付薪酬", "应付职工"],
    "应交税费": ["应交税费", "应交税"],
    "实收资本": ["实收资本", "实收"],
    "营业外收入": ["营业外收入"],
    "营业外支出": ["营业外支出"],
    "销售费用": ["销售费用", "销售费"],
    "管理费用": ["管理费用", "管理费"],
    "研发费用": ["研发费用", "研发费"],
    "财务费用": ["财务费用", "财务费"],
    "营业收入": ["营业收入", "营业收"],
    "营业成本": ["营业成本", "营业成"],
    "信用减值损失": ["信用减值损失", "信用减值"],
    "资产减值损失": ["资产减值损失", "资产减值"],
    "递延所得税资产": ["递延所得税资产"],
    "递延所得税负债": ["递延所得税负债"],
    "资本公积": ["资本公积"],
    "盈余公积": ["盈余公积"],
}

# 表类型后缀
TABLE_NAME_SUFFIXES = [
    "-账龄",
    "-组合",
    "-明细",
    "-分类",
    "-补充资料",
    "（续）",
    "(续)",
    "-按支付对象分类",
    "-按债务人分类",
]

# 占位符模式
PLACEHOLDER_PATTERNS = [
    re.compile(r"X{2,}"),  # XX, XXX, XXXX, XXXXX…
    re.compile(r"【[^】]*XX[^】]*】"),  # 【…XX…】
    re.compile(r"[（(][^）)]*待[^）)]*[）)]"),  # （待补充）、（待确认）
]

# 章节序号模式 — 表内小节标题（如"一、账面原值合计"）
SECTION_PATTERNS = [
    re.compile(r"^[一二三四五六七八九十]+[、．\.]"),  # 一、 一． 十、
    re.compile(r"^[（(][一二三四五六七八九十]+[）)]"),  # （一）（二）
    re.compile(r"^[（(]\d+[）)]"),  # （1）（2）
    re.compile(r"^\d+[．\.]、?"),  # 1． 2. 3、
    re.compile(r"^[(（]\d+[)）]"),  # (1)（2）
]

# 合计行关键词
TOTAL_KEYWORDS = ["合计", "小计", "总计"]
SUBTOTAL_KEYWORDS = ["小计"]

# 最大表名段落长度字符数
MAX_NAME_PARAGRAPH_LENGTH = 30


# ============================================================
# 模板分析器
# ============================================================


class TemplateAnalyzer:
    """
    Word 审计附注模板表格结构分析器。

    读取 .docx 文件，提取每个表格的结构信息：
    - 表名推断（基于前文段落）
    - 行结构（行名、合计/小计标记）
    - 列结构（表头）
    - 小节标题（一、二、…）
    - 占位符单元格
    """

    def __init__(self):
        self._doc = None
        self._doc_path = ""

    # ── 入口方法 ──────────────────────────────────────────

    def analyze(self, docx_path: str) -> Dict[str, Any]:
        """
        分析指定 Word 模板，返回结构化结果。

        参数:
            docx_path: .docx 文件路径

        返回:
            dict: 包含完整表格结构信息的字典

        异常:
            FileNotFoundError: 文件不存在
            ValueError: 文件无法解析
        """
        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"模板文件不存在: {docx_path}")
        if path.suffix.lower() not in (".docx",):
            raise ValueError(f"不支持的文件格式: {path.suffix}，仅支持 .docx")

        try:
            self._doc = Document(str(path))
            self._doc_path = str(path.resolve())
        except Exception as e:
            raise ValueError(f"无法读取 Word 文档: {e}")

        tables = self.analyze_all_tables(self._doc)

        result: Dict[str, Any] = {
            "file": self._doc_path,
            "file_name": path.name,
            "num_tables": len(tables),
            "tables": tables,
        }

        return result

    # ── 表格分析 ──────────────────────────────────────────

    def analyze_all_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """
        分析文档中所有表格。

        参数:
            doc: python-docx Document 对象

        返回:
            list[dict]: 每个表格的结构信息
        """
        tables_info: List[Dict[str, Any]] = []

        for idx, table in enumerate(doc.tables):
            try:
                table_info = self._analyze_single_table(doc, table, idx)
                tables_info.append(table_info)
            except Exception as e:
                # 单个表分析失败不影响后续表
                tables_info.append(
                    {
                        "index": idx,
                        "error": str(e),
                        "name": "",
                        "confidence": 0.0,
                        "num_rows": len(table.rows),
                        "num_cols": len(table.columns) if table.columns else 0,
                        "headers": [],
                        "section_headers": [],
                        "rows": [],
                        "total_rows": [],
                        "subtotal_rows": [],
                        "placeholder_cells": [],
                        "summary_types": [],
                    }
                )

        return tables_info

    def _analyze_single_table(self, doc: Document, table, index: int) -> Dict[str, Any]:
        """
        分析单个表格的结构。

        参数:
            doc: Document 对象（用于前文查找）
            table: python-docx Table 对象
            index: 表格在文档中的索引位置

        返回:
            dict: 该表格的结构信息
        """
        rows_data = self._extract_rows(table)
        headers = self._extract_headers(table)
        num_cols = self._count_columns(table)
        num_rows = len(table.rows)

        name, confidence = self.detect_table_name(doc, index)
        section_headers = self.detect_section_headers(table)
        total_rows = self.detect_total_rows(rows_data)
        subtotal_rows = self.detect_subtotal_rows(rows_data)
        placeholder_cells = self.detect_placeholders(table)

        # 合计行索引列表（用于快速查找）
        total_indices = [r["index"] for r in total_rows]
        subtotal_indices = [r["index"] for r in subtotal_rows]

        # 汇总表中出现的合计类型
        summary_types = sorted(
            set(r["name"].strip() for r in total_rows if r["name"].strip())
        )

        return {
            "index": index,
            "name": name,
            "confidence": round(confidence, 2),
            "num_rows": num_rows,
            "num_cols": num_cols,
            "headers": headers,
            "section_headers": section_headers,
            "rows": rows_data,
            "total_rows": total_indices,
            "subtotal_rows": subtotal_indices,
            "placeholder_cells": placeholder_cells,
            "summary_types": summary_types,
        }

    # ── 表名推断 ──────────────────────────────────────────

    def detect_table_name(self, doc: Document, table_index: int) -> tuple:
        """
        根据表格前的段落文本推断表格名称。

        策略:
          1. 在文档段落列表中定位该表格的位置
          2. 向前查找非空段落
          3. 跳过表头/标题样式段落
          4. 取最近且短于 MAX_NAME_PARAGRAPH_LENGTH 的段落
          5. 与 COMMON_NOTE_TABLE_NAMES 匹配

        参数:
            doc: Document 对象
            table_index: 表格索引

        返回:
            (name: str, confidence: float): 推断的表名和置信度
        """
        # 收集所有表格在段落流中的位置
        # python-docx 的 document.tables 按文档顺序排列
        # 段落和表格在 body.xml 中交错出现
        body = doc.element.body
        table_elements = body.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
        )
        para_elements = body.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
        )

        # 构建段落中的表格位置映射
        # 遍历 body 子元素，记录每个表格前最近的段落
        table_positions = []
        para_idx = 0
        tbl_counter = 0

        tbl_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
        p_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"

        for child in body:
            if child.tag == tbl_tag:
                table_positions.append(
                    {
                        "table_index": tbl_counter,
                        "preceding_paragraphs_end": para_idx - 1,
                    }
                )
                tbl_counter += 1
            elif child.tag == p_tag:
                para_idx += 1

        # 获取当前表格前最后一个段落的索引
        preceding_end = -1
        for tp in table_positions:
            if tp["table_index"] == table_index:
                preceding_end = tp["preceding_paragraphs_end"]
                break

        # 从该位置向前搜索表名
        candidates = []
        for i in range(preceding_end, -1, -1):
            try:
                para = doc.paragraphs[i]
            except IndexError:
                break

            text = para.text.strip()
            if not text:
                continue

            # 跳过明显不是表名的段落（长段落、编号段落、纯数字等）
            if len(text) > MAX_NAME_PARAGRAPH_LENGTH:
                candidates.append((text, 0.3))
                continue

            # 直接匹配标准附注表名
            matched_name, score = self._match_note_table_name(text)
            if matched_name:
                return matched_name, score

            # 匹配关键词
            matched_name, score = self._match_keywords(text)
            if matched_name:
                return matched_name, score

            # 短文本且有中文字符 → 作为低置信度候选
            if re.search(r"[\u4e00-\u9fff]", text):
                candidates.append((text, 0.4))

        # 没有直接命中的候选者
        if candidates:
            best = max(candidates, key=lambda x: x[1])
            return best[0], best[1]

        return "", 0.0

    def _match_note_table_name(self, text: str) -> tuple:
        """
        尝试将文本与标准附注表名精确匹配。

        参数:
            text: 待匹配文本

        返回:
            (name, score) 或 ("", 0.0)
        """
        # 去除无关字符后比较
        clean = re.sub(r"\s+", "", text)
        clean = clean.replace("　", "")

        for name in COMMON_NOTE_TABLE_NAMES:
            # 精确相等
            if clean == name.replace(" ", ""):
                return name, 0.95

            # 表名前缀匹配（如 "货币资金" 匹配 "货币资金-明细"）
            for suffix in TABLE_NAME_SUFFIXES:
                full_name = f"{name}{suffix}"
                if clean == full_name.replace(" ", ""):
                    return full_name, 0.9

            # 包含关系
            norm_name = name.replace(" ", "")
            if len(norm_name) >= 3 and norm_name in clean:
                return name, 0.7

        return "", 0.0

    def _match_keywords(self, text: str) -> tuple:
        """
        通过关键词匹配表名。

        参数:
            text: 待匹配文本

        返回:
            (name, score) 或 ("", 0.0)
        """
        clean = re.sub(r"\s+", "", text)
        clean = clean.replace("　", "")

        for name, keywords in TABLE_NAME_KEYWORDS.items():
            for kw in keywords:
                if kw in clean:
                    return name, 0.65

        return "", 0.0

    # ── 行数据提取 ───────────────────────────────────────

    def _extract_rows(self, table) -> List[Dict[str, Any]]:
        """
        提取表格所有行的数据。

        参数:
            table: python-docx Table 对象

        返回:
            list[dict]: 行数据列表
        """
        rows_data: List[Dict[str, Any]] = []

        for ri, row in enumerate(table.rows):
            cells = row.cells
            if not cells:
                continue

            # 第一格 → 行名（trimmed）
            row_name = cells[0].text.strip() if len(cells) > 0 else ""

            # 其余格 → 值
            values = [
                self._clean_cell_text(cells[ci].text) for ci in range(1, len(cells))
            ]

            is_total = self._is_total_row(row_name)
            is_subtotal = self._is_subtotal_row(row_name)

            rows_data.append(
                {
                    "index": ri,
                    "name": row_name,
                    "is_total": is_total,
                    "is_subtotal": is_subtotal,
                    "values": values,
                }
            )

        return rows_data

    def _extract_headers(self, table) -> List[str]:
        """
        提取表头（第一行列文本）。

        参数:
            table: python-docx Table 对象

        返回:
            list[str]: 列头文本列表
        """
        if not table.rows:
            return []

        first_row = table.rows[0]
        headers = []
        for cell in first_row.cells:
            headers.append(self._clean_cell_text(cell.text))

        return headers

    def _count_columns(self, table) -> int:
        """
        统计表格列数（取最大行列数）。

        参数:
            table: python-docx Table 对象

        返回:
            int: 列数
        """
        max_cols = 0
        for row in table.rows:
            if len(row.cells) > max_cols:
                max_cols = len(row.cells)
        return max_cols

    # ── 小节标题检测 ─────────────────────────────────────

    def detect_section_headers(self, table) -> List[str]:
        """
        检测表内小节标题行。

        典型模式:
          - "一、账面原值"
          - "（一）房屋建筑物"
          - "1．房屋及建筑物"
          - "2. 机器设备"

        参数:
            table: python-docx Table 对象

        返回:
            list[str]: 检测到的小节标题文本
        """
        section_headers: List[str] = []

        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not text:
                    continue
                for pattern in SECTION_PATTERNS:
                    if pattern.search(text):
                        if text not in section_headers:
                            section_headers.append(text)
                        break

        return section_headers

    # ── 合计行检测 ───────────────────────────────────────

    def detect_total_rows(self, rows: List[Dict]) -> List[Dict]:
        """
        检测合计行（含"合计""总计"等关键词）。

        参数:
            rows: _extract_rows 的输出

        返回:
            list[dict]: 合计行的行数据
        """
        return [r for r in rows if r["is_total"] and not r["is_subtotal"]]

    def detect_subtotal_rows(self, rows: List[Dict]) -> List[Dict]:
        """
        检测小计行（含"小计"关键词）。

        参数:
            rows: _extract_rows 的输出

        返回:
            list[dict]: 小计行的行数据
        """
        return [r for r in rows if r["is_subtotal"]]

    def _is_total_row(self, row_name: str) -> bool:
        """判断行名是否为合计行。"""
        if not row_name:
            return False
        clean = re.sub(r"\s+", "", row_name)
        for kw in TOTAL_KEYWORDS:
            if kw in clean:
                return True
        return False

    def _is_subtotal_row(self, row_name: str) -> bool:
        """判断行名是否为小计行。"""
        if not row_name:
            return False
        clean = re.sub(r"\s+", "", row_name)
        for kw in SUBTOTAL_KEYWORDS:
            if kw in clean:
                return True
        return False

    # ── 占位符检测 ───────────────────────────────────────

    def detect_placeholders(self, table) -> List[Dict[str, Any]]:
        """
        检测表格中的占位符单元格。

        检测模式:
          - XX, XXX, XXXX（连续大写X）
          - 【…XX…】（方括号含XX）
          - （待补充）（待确认）

        参数:
            table: python-docx Table 对象

        返回:
            list[dict]: 占位符单元格列表
                [{ "row": int, "col": int, "text": str }, …]
        """
        placeholders: List[Dict[str, Any]] = []

        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                for pattern in PLACEHOLDER_PATTERNS:
                    if pattern.search(text):
                        # 避免重复（同一格可能匹配多个模式）
                        if not any(
                            p["row"] == ri and p["col"] == ci for p in placeholders
                        ):
                            placeholders.append(
                                {
                                    "row": ri,
                                    "col": ci,
                                    "text": text,
                                }
                            )
                        break

        return placeholders

    # ── 辅助方法 ─────────────────────────────────────────

    @staticmethod
    def _clean_cell_text(text: str) -> str:
        """清理单元格文本（去除多余空白、不可见字符）。"""
        if not text:
            return ""
        text = re.sub(r"\s+", " ", text)  # 多个空格→单个
        text = text.replace("\u3000", "")  # 全角空格
        text = text.replace("\n", " ").replace("\r", "")
        return text.strip()

    # ── JSON 序列化 ──────────────────────────────────────

    @staticmethod
    def to_json(
        result: Dict[str, Any], ensure_ascii: bool = False, indent: int = 2
    ) -> str:
        """
        将分析结果序列化为 JSON。

        参数:
            result: analyze() 的返回值
            ensure_ascii: 是否转义非ASCII字符（默认 False 保留中文）
            indent: 缩进空格数

        返回:
            str: JSON 字符串
        """
        return json.dumps(result, ensure_ascii=ensure_ascii, indent=indent, default=str)


# ============================================================
# 简易统计报告
# ============================================================


def generate_summary(result: Dict[str, Any]) -> str:
    """
    生成分析结果的摘要文本报告。

    参数:
        result: analyze() 的返回值

    返回:
        str: 格式化的文本报告
    """
    lines = []
    lines.append(f"文件: {result['file_name']}")
    lines.append(f"表格总数: {result['num_tables']}")
    lines.append("")

    for tbl in result["tables"]:
        if tbl.get("error"):
            lines.append(f"  [表 {tbl['index']}] 错误: {tbl['error']}")
            continue

        name_display = tbl["name"] if tbl["name"] else "(未识别)"
        confidence_display = f" (置信度: {tbl['confidence']})" if tbl["name"] else ""

        lines.append(f"  [表 {tbl['index']}] {name_display}{confidence_display}")
        lines.append(f"      行×列: {tbl['num_rows']}×{tbl['num_cols']}")
        lines.append(
            f"      表头: {tbl['headers'][:6]}{'…' if len(tbl['headers']) > 6 else ''}"
        )

        if tbl["section_headers"]:
            lines.append(
                f"      小节标题 ({len(tbl['section_headers'])}): "
                f"{tbl['section_headers'][:4]}"
                f"{'…' if len(tbl['section_headers']) > 4 else ''}"
            )

        if tbl["total_rows"]:
            lines.append(f"      合计行: {tbl['total_rows']}")
        if tbl["subtotal_rows"]:
            lines.append(f"      小计行: {tbl['subtotal_rows']}")

        if tbl["placeholder_cells"]:
            lines.append(f"      占位符: {len(tbl['placeholder_cells'])} 处")
            # 列出前5个占位符
            for pc in tbl["placeholder_cells"][:5]:
                lines.append(
                    f"        第{pc['row'] + 1}行第{pc['col'] + 1}列: "
                    f"「{pc['text'][:30]}」"
                )
            if len(tbl["placeholder_cells"]) > 5:
                lines.append(f"        … 还有 {len(tbl['placeholder_cells']) - 5} 处")

        lines.append("")

    # 总体统计
    total_placeholders = sum(
        len(t.get("placeholder_cells", []))
        for t in result["tables"]
        if "error" not in t
    )
    total_sections = sum(
        len(t.get("section_headers", [])) for t in result["tables"] if "error" not in t
    )
    named_tables = sum(
        1 for t in result["tables"] if t.get("name") and "error" not in t
    )

    lines.append(f"汇总:")
    lines.append(f"  已识别表名: {named_tables}/{result['num_tables']}")
    lines.append(f"  小节标题总数: {total_sections}")
    lines.append(f"  占位符总数: {total_placeholders}")

    return "\n".join(lines)


# ============================================================
# 独立入口
# ============================================================


def main():
    """CLI 入口：分析一个或多个 .docx 模板。"""
    import argparse

    parser = argparse.ArgumentParser(
        description="审计附注模板分析器 — 提取 .docx 中所有表格的结构信息",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python template_analyzer.py 模板.docx
  python template_analyzer.py --json 模板.docx > analysis.json
  python template_analyzer.py --summary-only 模板.docx
  python template_analyzer.py 模板1.docx 模板2.docx
        """,
    )
    parser.add_argument("files", nargs="+", help="要分析的 .docx 模板文件路径")
    parser.add_argument("--json", "-j", action="store_true", help="以 JSON 格式输出")
    parser.add_argument(
        "--summary-only",
        "-s",
        action="store_true",
        help="仅输出摘要文本（不输出完整JSON）",
    )
    parser.add_argument(
        "--indent", "-i", type=int, default=2, help="JSON 缩进空格数（默认 2）"
    )
    parser.add_argument(
        "--table-filter",
        "-t",
        type=int,
        nargs="+",
        help="只分析指定的表索引（如 -t 0 1 3）",
    )

    args = parser.parse_args()

    analyzer = TemplateAnalyzer()
    results = []

    for filepath in args.files:
        path = Path(filepath)
        if not path.exists():
            print(f"跳过: 文件不存在 — {filepath}", file=sys.stderr)
            continue

        try:
            result = analyzer.analyze(filepath)

            # 表格筛选
            if args.table_filter is not None:
                result["tables"] = [
                    t for t in result["tables"] if t["index"] in args.table_filter
                ]
                result["num_tables"] = len(result["tables"])

            results.append(result)
        except Exception as e:
            print(f"错误: {filepath} — {e}", file=sys.stderr)
            continue

    if not results:
        sys.exit(1)

    if args.json:
        # JSON 输出
        if len(results) == 1:
            output = TemplateAnalyzer.to_json(results[0], indent=args.indent)
        else:
            output = TemplateAnalyzer.to_json(results, indent=args.indent)
        sys.stdout.write(output)
        sys.stdout.write("\n")
    elif args.summary_only:
        # 仅摘要
        for r in results:
            print(generate_summary(r))
            print("=" * 60)
    else:
        # 完整文本输出
        for r in results:
            print(generate_summary(r))
            # 输出完整结构化信息
            print("\n完整结构:")
            print(TemplateAnalyzer.to_json(r, indent=args.indent))
            print("=" * 60)


if __name__ == "__main__":
    main()
